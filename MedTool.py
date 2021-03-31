#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import datetime,  sys,PyQt5, openpyxl, itertools
import pandas as pd
from docx import Document
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QApplication


# In[ ]:


starttime = str(datetime.datetime.now())
datetoday = datetime.date.today()
datetoday = datetoday.strftime("%m/%d/%y")
print(f'Mediation started on {datetoday}!')
demands = []
offers = []
offertime = []
demandtime = []
mediationinfo = []
bracketHigh = []
bracketLow = []
midPoints = []
memoList=[]
a = []
data = {"Demands": demands,
        "Demand_Time": demandtime,
        "Offers": offers,
        "Offer_Time": offertime,
        "Mid_Point": midPoints}

counter = 0


# In[ ]:


class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.setWindowModality(QtCore.Qt.WindowModal)
        MainWindow.resize(863, 772)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(MainWindow.sizePolicy().hasHeightForWidth())
        MainWindow.setSizePolicy(sizePolicy)
        MainWindow.setMinimumSize(QtCore.QSize(613, 491))
        MainWindow.setMaximumSize(QtCore.QSize(1700, 1500))
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.verticalLayout_4 = QtWidgets.QVBoxLayout()
        self.verticalLayout_4.setContentsMargins(-1, -1, -1, 0)
        self.verticalLayout_4.setObjectName("verticalLayout_4")
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_3.setSizeConstraint(QtWidgets.QLayout.SetNoConstraint)
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        self.formLayout = QtWidgets.QFormLayout()
        self.formLayout.setSizeConstraint(QtWidgets.QLayout.SetNoConstraint)
        self.formLayout.setObjectName("formLayout")
        self.cASENUMBERLabel = QtWidgets.QLabel(self.centralwidget)
        self.cASENUMBERLabel.setFrameShape(QtWidgets.QFrame.WinPanel)
        self.cASENUMBERLabel.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.cASENUMBERLabel.setObjectName("cASENUMBERLabel")
        self.formLayout.setWidget(0, QtWidgets.QFormLayout.LabelRole, self.cASENUMBERLabel)
        self.cASENUMBERLineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.cASENUMBERLineEdit.setObjectName("cASENUMBERLineEdit")
        self.formLayout.setWidget(0, QtWidgets.QFormLayout.FieldRole, self.cASENUMBERLineEdit)
        self.mEDIATORLabel = QtWidgets.QLabel(self.centralwidget)
        self.mEDIATORLabel.setFrameShape(QtWidgets.QFrame.WinPanel)
        self.mEDIATORLabel.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.mEDIATORLabel.setObjectName("mEDIATORLabel")
        self.formLayout.setWidget(1, QtWidgets.QFormLayout.LabelRole, self.mEDIATORLabel)
        self.mEDIATORLineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.mEDIATORLineEdit.setObjectName("mEDIATORLineEdit")
        self.formLayout.setWidget(1, QtWidgets.QFormLayout.FieldRole, self.mEDIATORLineEdit)
        self.pLAINTIFFSATTORNEYLabel = QtWidgets.QLabel(self.centralwidget)
        self.pLAINTIFFSATTORNEYLabel.setFrameShape(QtWidgets.QFrame.WinPanel)
        self.pLAINTIFFSATTORNEYLabel.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.pLAINTIFFSATTORNEYLabel.setObjectName("pLAINTIFFSATTORNEYLabel")
        self.formLayout.setWidget(2, QtWidgets.QFormLayout.LabelRole, self.pLAINTIFFSATTORNEYLabel)
        self.pLAINTIFFSATTORNEYLineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.pLAINTIFFSATTORNEYLineEdit.setObjectName("pLAINTIFFSATTORNEYLineEdit")
        self.formLayout.setWidget(2, QtWidgets.QFormLayout.FieldRole, self.pLAINTIFFSATTORNEYLineEdit)
        self.pushButtonForm = QtWidgets.QPushButton(self.centralwidget)
        self.pushButtonForm.setMinimumSize(QtCore.QSize(3, 0))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.pushButtonForm.setFont(font)
        self.pushButtonForm.setObjectName("pushButtonForm")
        self.formLayout.setWidget(3, QtWidgets.QFormLayout.FieldRole, self.pushButtonForm)
        self.pushButtonSave = QtWidgets.QPushButton(self.centralwidget)
        self.pushButtonSave.setObjectName("pushButtonSave")
        self.formLayout.setWidget(4, QtWidgets.QFormLayout.FieldRole, self.pushButtonSave)
        self.horizontalLayout_3.addLayout(self.formLayout)
        spacerItem = QtWidgets.QSpacerItem(10, 61, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.horizontalLayout_3.addItem(spacerItem)
        self.listWidget = QtWidgets.QListWidget(self.centralwidget)
        self.listWidget.setObjectName("listWidget")
        self.horizontalLayout_3.addWidget(self.listWidget)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.horizontalLayout_3.addItem(spacerItem1)
        self.verticalLayout_4.addLayout(self.horizontalLayout_3)
        spacerItem2 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.verticalLayout_4.addItem(spacerItem2)
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label.setFont(font)
        self.label.setAlignment(QtCore.Qt.AlignHCenter|QtCore.Qt.AlignTop)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label, 0, QtCore.Qt.AlignHCenter|QtCore.Qt.AlignVCenter)
        self.lineDemand1 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineDemand1.setObjectName("lineDemand1")
        self.verticalLayout.addWidget(self.lineDemand1)
        self.pushButtonDemand1 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButtonDemand1.setObjectName("pushButtonDemand1")
        self.verticalLayout.addWidget(self.pushButtonDemand1)
        self.horizontalLayout_2.addLayout(self.verticalLayout)
        spacerItem3 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.horizontalLayout_2.addItem(spacerItem3)
        self.verticalLayout_2 = QtWidgets.QVBoxLayout()
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_2.setFont(font)
        self.label_2.setAlignment(QtCore.Qt.AlignCenter)
        self.label_2.setObjectName("label_2")
        self.verticalLayout_2.addWidget(self.label_2, 0, QtCore.Qt.AlignHCenter|QtCore.Qt.AlignVCenter)
        self.lineOffer1 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineOffer1.setObjectName("lineOffer1")
        self.verticalLayout_2.addWidget(self.lineOffer1)
        self.pushButtonOffer1 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButtonOffer1.setObjectName("pushButtonOffer1")
        self.verticalLayout_2.addWidget(self.pushButtonOffer1)
        self.horizontalLayout_2.addLayout(self.verticalLayout_2)
        self.verticalLayout_4.addLayout(self.horizontalLayout_2)
        self.pushButtonBracket = QtWidgets.QPushButton(self.centralwidget)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Maximum, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.pushButtonBracket.sizePolicy().hasHeightForWidth())
        self.pushButtonBracket.setSizePolicy(sizePolicy)
        self.pushButtonBracket.setObjectName("pushButtonBracket")
        self.verticalLayout_4.addWidget(self.pushButtonBracket, 0, QtCore.Qt.AlignHCenter)
        spacerItem4 = QtWidgets.QSpacerItem(40, 50, QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Minimum)
        self.verticalLayout_4.addItem(spacerItem4)
        self.tableWidget = QtWidgets.QTableWidget(self.centralwidget)
        self.tableWidget.setEnabled(True)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.MinimumExpanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.tableWidget.sizePolicy().hasHeightForWidth())
        self.tableWidget.setSizePolicy(sizePolicy)
        self.tableWidget.setMaximumSize(QtCore.QSize(16777215, 16777215))
        self.tableWidget.setAlternatingRowColors(True)
        self.tableWidget.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tableWidget.setTextElideMode(QtCore.Qt.ElideLeft)
        self.tableWidget.setColumnCount(4)
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setHorizontalHeaderLabels(['Type','Amount','Time','Mid-Point'])
        self.tableWidget.setRowCount(0)
        self.verticalLayout_4.addWidget(self.tableWidget)
        self.horizontalLayout_5 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_5.setObjectName("horizontalLayout_5")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.horizontalLayout_5.addWidget(self.label_3)
        self.textBrowserNotes = QtWidgets.QTextBrowser(self.centralwidget)
        self.textBrowserNotes.setTextInteractionFlags(QtCore.Qt.LinksAccessibleByKeyboard|QtCore.Qt.LinksAccessibleByMouse|QtCore.Qt.TextBrowserInteraction|QtCore.Qt.TextEditable|QtCore.Qt.TextSelectableByMouse)
        self.textBrowserNotes.setObjectName("textBrowserNotes")
        self.horizontalLayout_5.addWidget(self.textBrowserNotes)
        spacerItem5 = QtWidgets.QSpacerItem(40, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.horizontalLayout_5.addItem(spacerItem5)
        self.checkBox = QtWidgets.QCheckBox(self.centralwidget)
        self.checkBox.setObjectName("checkBox")
        self.horizontalLayout_5.addWidget(self.checkBox)
        self.buttonBox = QtWidgets.QDialogButtonBox(self.centralwidget)
        self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.Cancel|QtWidgets.QDialogButtonBox.Ok)
        self.buttonBox.setObjectName("buttonBox")
        self.horizontalLayout_5.addWidget(self.buttonBox)
        self.verticalLayout_4.addLayout(self.horizontalLayout_5)
        self.gridLayout.addLayout(self.verticalLayout_4, 0, 0, 1, 1)
        MainWindow.setCentralWidget(self.centralwidget)
        self.actionHow_to_Use = QtWidgets.QAction(MainWindow)


        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        
        self.onlyInt = QtGui.QIntValidator()
        self.lineDemand1.setValidator(self.onlyInt)
        self.lineOffer1.setValidator(self.onlyInt)
        self.retranslateUi(MainWindow)
        self.pushButtonSave.clicked.connect(self.mediation_info)
        self.pushButtonForm.clicked.connect(self.listWidget.clear)
        self.pushButtonForm.clicked.connect(self.pLAINTIFFSATTORNEYLineEdit.clear)
        self.pushButtonForm.clicked.connect(self.mEDIATORLineEdit.clear)
        self.pushButtonForm.clicked.connect(self.cASENUMBERLineEdit.clear)
        self.buttonBox.rejected.connect(self.tableWidget.clearContents)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.pushButtonForm.clicked.connect(self.listWidget.reset)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.buttonBox.rejected.connect(self.clear_data)
        self.buttonBox.rejected.connect(self.tableWidget.clearContents)
        self.buttonBox.accepted.connect(self.final_offers)
        self.pushButtonDemand1.clicked.connect(self.demand_update)
        self.pushButtonOffer1.clicked.connect(self.offer_update)
        self.pushButtonOffer1.clicked.connect(self.mid_point)
        self.pushButtonOffer1.clicked.connect(self.round_counter)
        self.pushButtonBracket.clicked.connect(self.set_bracket)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
      

    
    
    def mediation_info(self):
        '''grabs the mediation information from the
        form in the top left'''
        mediationinfo.append(self.cASENUMBERLineEdit.text())
        mediationinfo.append(self.mEDIATORLineEdit.text())
        mediationinfo.append(self.pLAINTIFFSATTORNEYLineEdit.text())
        
        if not self.cASENUMBERLineEdit.text() or not self.mEDIATORLineEdit or not self.pLAINTIFFSATTORNEYLineEdit:
            print("Please complete the form in the top left.")
        else:
            self.listWidget.addItem(self.cASENUMBERLineEdit.text())
            self.listWidget.addItem(self.mEDIATORLineEdit.text())
            self.listWidget.addItem(self.pLAINTIFFSATTORNEYLineEdit.text())
        
        
    def demand_update(self):
        '''updates the demand list of global variables.
        adds the dollar amount and time the demand was made to the demand list
        adds demand to list in order that the demand was made
        will use the form data to fill in demand argument'''
        try:
            demtime = str(datetime.datetime.now())
            dem = int(self.lineDemand1.text())
            print(f'The current demand is {dem} at {demtime.split()[1]}')
            demands.append(dem)
            demandtime.append(demtime.split()[1])
            rowPosition = self.tableWidget.rowCount()
            self.tableWidget.insertRow(rowPosition)
            self.tableWidget.setItem(rowPosition,0, QtWidgets.QTableWidgetItem('Demand'))
            self.tableWidget.setItem(rowPosition,1, QtWidgets.QTableWidgetItem(f'${dem}'))
            self.tableWidget.setItem(rowPosition,2, QtWidgets.QTableWidgetItem(demtime.split()[1]))
        except ValueError:
            print("Please input a demand!")
        

    def offer_update(self):
        '''updates the offer list of global variables
        adds the dollar amount and time the offer was made to the offer list
        adds offer to list in order that offer was made
        will use the form data to fill in offer argument'''
        try:
            offtime = str(datetime.datetime.now())
            off = int(self.lineOffer1.text())
            print(f'The current offer is {off} at {offtime.split()[1]}')
            offers.append(off)
            memoList.append(tuple(["Demand and offer",demands[-1], off]))
            offertime.append(offtime.split()[1])
            rowPosition = self.tableWidget.rowCount()
            self.tableWidget.insertRow(rowPosition)
            self.tableWidget.setItem(rowPosition,0, QtWidgets.QTableWidgetItem('Offer'))
            self.tableWidget.setItem(rowPosition,1, QtWidgets.QTableWidgetItem(f'${off}'))
            self.tableWidget.setItem(rowPosition,2, QtWidgets.QTableWidgetItem(offtime.split()[1]))
        except ValueError:
            print("Please input an offer!")
    
    def set_bracket(self):
        '''
        updates the bracket global lists
        outputs them to the table and lists it as bracket
        '''
        try:
            
            bracketTime = str(datetime.datetime.now())
            bracket1 = int(self.lineDemand1.text())
            bracket2 = int(self.lineOffer1.text())
            memoList.append(tuple(["Bracket", bracket1, bracket2]))
            midPoint = (bracket1 + bracket2)/2
            print(f'The current bracket is {bracket1} / {bracket2} at {bracketTime.split()[1]}')
            bracketHigh.append(bracket1)
            bracketLow.append(bracket2)
            rowPosition = self.tableWidget.rowCount()
            self.tableWidget.insertRow(rowPosition)
            self.tableWidget.setItem(rowPosition,0, QtWidgets.QTableWidgetItem('Bracket'))
            self.tableWidget.setItem(rowPosition,1, QtWidgets.QTableWidgetItem(f'${bracket1} / ${bracket2}'))
            self.tableWidget.setItem(rowPosition,2, QtWidgets.QTableWidgetItem(bracketTime.split()[1]))
            self.tableWidget.setItem(rowPosition,3, QtWidgets.QTableWidgetItem(f'${midPoint}'))
        except ValueError:
            print("Please input a High Bracket and a Low Bracket")
    
    def mid_point(self):
        '''Will calculate the mid point between the latest demand and
        the later offer'''
        midTime = str(datetime.datetime.now())
        try:
            midPoint = (offers[-1] + demands[-1])/2
            midPoints.append(midPoint)
            rowPosition = self.tableWidget.rowCount()
            self.tableWidget.insertRow(rowPosition)
            self.tableWidget.setItem(rowPosition,0, QtWidgets.QTableWidgetItem('MidPoint'))
            self.tableWidget.setItem(rowPosition,1, QtWidgets.QTableWidgetItem('N/A'))
            self.tableWidget.setItem(rowPosition,2, QtWidgets.QTableWidgetItem(midTime.split()[1]))
            self.tableWidget.setItem(rowPosition,3, QtWidgets.QTableWidgetItem(f'${midPoint}'))
        except IndexError:
            print("Please input an offer and a demand!")            
    
    def round_counter(self):
        '''Will keep track of the rounds of demands and offers
        uses the global variables for the counter'''
        global counter
        counter = counter +1
        rowPosition = self.tableWidget.rowCount()
        self.tableWidget.insertRow(rowPosition)
        self.tableWidget.setItem(rowPosition,0, QtWidgets.QTableWidgetItem(f'Round {counter}'))
        self.tableWidget.setItem(rowPosition,1, QtWidgets.QTableWidgetItem('N/A'))
        self.tableWidget.setItem(rowPosition,2, QtWidgets.QTableWidgetItem('N/A'))
        self.tableWidget.setItem(rowPosition,3, QtWidgets.QTableWidgetItem('N/A'))
    
        
    def final_offers(self): 
        '''
        Checks all fields are entered
        will save excel sheet to default directory
        will save a memo of the mediation to directory
        '''
        case = str(self.cASENUMBERLineEdit.text())
        mediator = str(self.mEDIATORLineEdit.text())
        pltatt = str(self.pLAINTIFFSATTORNEYLineEdit.text())
        sheet_name = f'Mediation {mediator} and {pltatt}'
        sheet_info = sheet_name[:30] #sheet name cannot be longer than 31 characters
        document = Document()
        document.add_heading(f'Mediation Summary {self.cASENUMBERLineEdit.text()}', 0)
        if not case or not mediator or not pltatt:
            print("Please input the case number, mediator, or Plaintiff's Attorney")
        else:
            try:
                medsummary = f'On {datetoday} we had the mediation with {self.mEDIATORLineEdit.text()} and {self.pLAINTIFFSATTORNEYLineEdit.text()}. During the mediation, we made {len(memoList)} moves. The initial demand was ${demands[0]} and ended with ${demands[-1]}. The initial offer was ${offers[0]} and the final offer was ${offers[-1]}. Any demands, offers, or brackets are as follows: '
                document.add_paragraph(medsummary)
                for a in memoList:
                    document.add_paragraph(f'{a}', style='List Number')
                if self.checkBox.isChecked() == True:
                    document.add_paragraph('Mediation was successful. As such, we will draft closing documents and prepare to close the file.')
                else:
                    document.add_paragraph("Mediation was unsuccessful. As such, we should work to investigate the case further and plan the next course of action.")
                document.add_paragraph(self.textBrowserNotes.toPlainText())
                document.save(f'{case}_Mediation_Memo.docx')
                a = list(map(list,itertools.zip_longest(data['Demands'], data['Demand_Time'],data['Offers'], data['Offer_Time'], data['Mid_Point'])))
                df1 = pd.DataFrame()
                df1[['Demands','Demand_Time','Offers', "Offer_Time", "Mid_Points"]] = pd.DataFrame(a)
                df1[['Start_Time']] = starttime
                df1[['End_Time']] = str(datetime.datetime.now())
                df1.to_excel(case +".xlsx", sheet_name = sheet_info)
                sys.exit(app.exec_())
            except IndexError:
                print("Please input at least one demand and one offer")
            except OSError:
                print('Please make sure the case name does not include the following: ?, ", *, :')
                pass
                

    
    def clear_data(self):
        '''sets the list variables to global and allows them to be cleared if mistakes were input'''
        global demands, offers, offertime, demandtime, data,df1, counter, mediationinfo, bracketHigh, bracketLow, midPoints,memoList
        self.tableWidget.clearContents()
        self.tableWidget.setRowCount(0)
        demands = []
        offers = []
        offertime = []
        demandtime = []
        mediationinfo = []
        bracketHigh = []
        bracketLow = []
        midPoints = []
        memoList=[]
        data = {"Demands": demands,
                "Demand_Time": demandtime,
               "Offers": offers,
               "Offer_Time": offertime}
        df1 = pd.DataFrame()
        counter = 0
        print("All fields have been cleared")
        
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Mediation Tracker"))
        self.cASENUMBERLabel.setText(_translate("MainWindow", "CASE NUMBER"))
        self.mEDIATORLabel.setText(_translate("MainWindow", "MEDIATOR"))
        self.pLAINTIFFSATTORNEYLabel.setText(_translate("MainWindow", "PLAINTIFF\'S ATTORNEY"))
        self.pushButtonForm.setText(_translate("MainWindow", "Clear Form"))
        self.pushButtonSave.setText(_translate("MainWindow", "Save Information"))
        self.label.setText(_translate("MainWindow", "DEMAND/BRACKET HIGH"))
        self.pushButtonDemand1.setText(_translate("MainWindow", "Set Demand"))
        self.label_2.setText(_translate("MainWindow", "OFFER/BRACKET LOW"))
        self.pushButtonOffer1.setText(_translate("MainWindow", "Set Offer"))
        self.pushButtonBracket.setText(_translate("MainWindow", "Set Bracket"))
        self.label_3.setText(_translate("MainWindow", "Mediation Notes"))
        self.textBrowserNotes.setPlaceholderText(_translate("MainWindow", "Enter any mediation notes here(new medicals, information from counsel, any follow up for unsucessful mediations, etc...)"))
        self.checkBox.setText(_translate("MainWindow", "Mediation Successful?"))


# In[ ]:


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    ex = Ui_MainWindow()
    w = QtWidgets.QMainWindow()
    ex.setupUi(w)
    w.show()
    sys.exit(app.exec_())

